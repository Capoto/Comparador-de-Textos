from django.shortcuts import render

from django.http import HttpResponse
#from PyPDF2 import PdfReader
from PyPDF2 import PdfReader
import difflib
from spire.doc import *
from spire.doc.common import *
import requests
import xmltodict
import csv
import zipfile
import io
from django.conf import settings
from docx import Document
from docx.shared import Inches
from docx.shared import Pt,RGBColor
import pandas as pd


import warnings
warnings.filterwarnings("ignore",category=DeprecationWarning)

import datetime
import pandas as pd



def download_zip(l,nome):
    # create a zip file on-the-fly
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w') as zip_file:
        # add files to the zip file
        for i in l:
            zip_file.write(i)
        
    
    # create the HttpResponse object with the appropriate Content-Type header.
    response = HttpResponse(buffer.getvalue(), content_type='application/x-zip-compressed')
    # set the Content-Disposition header to force download of the file
    response['Content-Disposition'] = 'attachment;filename='+nome+".zip"
    return response 


def index(request):

    if request.method=="POST":
        form = request.FILES['myfile']
        form2 = request.FILES['myfile2']
        texto1=""
        texto2=""
        if  str(form).endswith('.pdf'):
            reader = PdfReader(form)
            
            with open('arquivo1.txt', 'w', encoding="utf-8") as f:
                for i in range(len(reader.pages)):
                    f.write(reader.pages[i].extract_text())
                    texto1+=reader.pages[i].extract_text()
        if  str(form2).endswith('.pdf'):
            reader = PdfReader(form2)
            with open('arquivo2.txt', 'w', encoding="utf-8") as f:
                for i in range(len(reader.pages)):
                    f.write(reader.pages[i].extract_text())
                    texto2+=reader.pages[i].extract_text()

        if  str(form).endswith('.docx'):
            # Create a Document object
            document = Document()
            # Load a Word document
            document.LoadFromFile(form)

            # Extract the text of the document
            document_text = document.GetText()
            texto1= document.GetText()
            # Write the extracted text into a text file
            with open("arquivo1.txt", "w", encoding="utf-8") as file:
                file.write(document_text)
        if  str(form2).endswith('.docx'):
            # Create a Document object
            document = Document()
            # Load a Word document
            document.LoadFromFile(str(form))

            # Extract the text of the document
            document_text = document.GetText()
            texto2= document.GetText()
            print(texto2)
            # Write the extracted text into a text file
            with open("arquivo1.txt", "w", encoding="utf-8") as file:
                file.write(document_text)
        
        d = difflib.HtmlDiff()
        html_diff = d.make_file(texto1.splitlines(), texto2.splitlines()) # a,b were defined earlier
        with open("diff.html", "w", encoding="utf-8") as f:
            f.write(html_diff)  
        
    return render(request,'index.html')


def plenario(request):
    lista = []
    if request.method=="POST":
        print("ok")
        data = request.POST
        action = data.get("plenariodata")
        ano = action.split('-')
        action = action.replace("-","")
        api_end_point = "https://legis.senado.leg.br/dadosabertos/plenario/lista/votacao/"+str(action)
        joke = requests.get(api_end_point)
        xpars = xmltodict.parse(joke.text)
        votos = xpars['ListaVotacoes']

        

        if 'Votacoes' not in votos:
            print("Não teve votação nesse dia")
            return render(request,'plenario.html',{'alerta': "verdadeiro"})

       

        else:
            

            votos = xpars['ListaVotacoes']['Votacoes']
            cont =1
            

            if type(votos['Votacao'])==list:
            
                for i in votos['Votacao']:
               
                    api_end_point2 = "https://legis.senado.leg.br/dadosabertos/materia/"+str(i['CodigoMateria'])
                    joke = requests.get(api_end_point2)
                    xpars = xmltodict.parse(joke.text)
                    materia =  xpars['DetalheMateria']
                    sim = 0
                    nao = 0
                    presidente  = 0 
                    abss = 0 
                    pnrv=0
                    ncom =0
                    desc = ""
                    mtr= "Sem Ementa"
                    res = ""
                    secreta=""
                    d= ""
                    l = [] 
                    orientada = []
                    id = ""
                

                    if 'Materia' in materia:
                        materia =  xpars['DetalheMateria']['Materia']['DadosBasicosMateria']['EmentaMateria']
                        mtr = materia
                        desc= i['DescricaoVotacao']
                        d = i['DataSessao']

                        id = i['DescricaoIdentificacaoMateria']
                        if i['Resultado']=="A":
                            res ="Aprovada"
                        else:
                            res = "Reprovada"

                        if i['Secreta']=="S":
                            secreta ="Sim"
                        else:
                            secreta = "Não"

                        for j in i['Votos']['VotoParlamentar']:
                            
                            if 'Voto' in j:
                                if j['Voto']=='Sim':
                                    sim+=1
                                elif j['Voto']=='Não':
                                    nao+=1
                                elif j['Voto']=='P-NRV':
                                    pnrv+=1
                                elif j['Voto']=='Presidente (art. 51 RISF)':
                                    presidente+=1
                                elif  j['Voto']=='Abstenção':
                                    abss+=1
                                elif j['Voto']=='NCom':
                                    ncom+=1
                        

                        
                                l.append([j['NomeParlamentar'],j['SiglaPartido'],j['SiglaUF'],j['Voto']])


                     
                   

                    else:

                        desc= i['DescricaoVotacao']
                        d = i['DataSessao']
                        id = i['DescricaoIdentificacaoMateria']
                        if i['Resultado']=="A":
                            res ="Aprovada"
                        else:
                            res = "Reprovada"

                        if i['Secreta']=="S":
                            secreta ="Sim"
                        else:
                            secreta = "Não"
                    
                        for j in i['Votos']['VotoParlamentar']:
                            if 'Voto' in j:
                            
                                if j['Voto']=='Sim':
                                    sim+=1
                                elif j['Voto']=='Não':
                                    nao+=1
                                elif j['Voto']=='P-NRV':
                                    pnrv+=1
                                elif j['Voto']=='Presidente (art. 51 RISF)':
                                    presidente+=1
                                elif  j['Voto']=='Abstenção':
                                    abss+=1
                                elif j['Voto']=='NCom':
                                    ncom+=1
                
                       
                                l.append([j['NomeParlamentar'],j['SiglaPartido'],j['SiglaUF'],j['Voto']])

                    api_end_point3 = "https://legis.senado.leg.br/dadosabertos/plenario/votacao/orientacaoBancada/"+action
                    joke = requests.get(api_end_point3)
                    xpars = xmltodict.parse(joke.text)
                    ori =  xpars['OrientacaoBancada']
                
                    if ori is None:
                        return render(request,'plenario.html')
                    
                    if 'votacoes':
                        ori = xpars['OrientacaoBancada']['votacoes']
                        flag = 0
                        for i  in  ori:
                            
                            if 'orientacoesLideranca' in i:
                                
                            
                                for j in i['orientacoesLideranca']:
                                    z = j['partido']
                                    if z=="Governo":
                                        print("teste")
                                        flag =1
                                        total = sim+nao+presidente+abss+pnrv+ncom
                                        orientada.append([d,desc,mtr,sim,nao,presidente,abss,pnrv,total,res,secreta, j['voto']])
                                        break
                                if flag==1:
                                    break
                        if flag==0:
                            total = sim+nao+presidente+abss+pnrv+ncom
                            orientada.append([d,desc,mtr,sim,nao,presidente,abss,pnrv,total,res,secreta,"SEM ORIENTAÇÃO"])


                    print(cont)
                    document = Document()
                    font = document.styles['Normal'].font
                    font.name = 'Arial'
                    font.size = Pt(12)
                    styles = document.styles['Heading 1'].font
                    styles.name = 'Arial'
                    paragraph_format = document.styles['Normal'].paragraph_format
                    paragraph_format.line_spacing = 1.15
                    z = open("CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv",'w+',encoding='cp1252')
               

                    document.add_heading('DataSessao ' + str(orientada[0][0]) +"\n", 0)

                    document.add_heading(id, 0)

                    document.add_paragraph("\n"+orientada[0][2])

                    document.add_paragraph("\n"+orientada[0][1])

                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Secreta'
                    hdr_cells[1].text = orientada[0][-2]

               
                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos'

                
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos Sim'
                    hdr_cells[1].text = str(orientada[0][3])

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos Não'
                    hdr_cells[1].text = str(orientada[0][4])
                
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Abstenção'
                    hdr_cells[1].text = str(orientada[0][6])


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos do Presidente'
                    hdr_cells[1].text = str(orientada[0][5])


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Não Compareceu'
                    hdr_cells[1].text = str(ncom)


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'PNRV'
                    hdr_cells[1].text = str(orientada[0][7])

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Total'
                    hdr_cells[1].text = str(orientada[0][8])


                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Resultado'
                    hdr_cells[1].text = orientada[0][-3]


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Orientação do Governo'
                    hdr_cells[1].text = orientada[0][-1]


                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Parlamentar'
                    hdr_cells[1].text = 'Partido'
                    hdr_cells[2].text = 'UF'
                    hdr_cells[3].text = 'Voto'

                
                    for iii in l:
                        table = document.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = iii[0]
                        hdr_cells[1].text =  iii[1]
                        hdr_cells[2].text = iii[2]
                        hdr_cells[3].text =  iii[3]
                



                    document.save('word/Votação '+str(cont)+ 'Plenário '+ano[2]+"_"+ano[1]+"_"+ano[0]+'.docx')
                    print(orientada)   
                    l.sort(key=lambda x:x[1])     
                    with open("CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv",'w', newline='',encoding='iso-8859-1') as csvfile:
                        writer = csv.writer(csvfile,delimiter =';')

                        # Write data for table 1
                        writer.writerow(["DataSessao",orientada[0][0]])
                      

                        # Add an empty line between tables
                        writer.writerow([])
                        
                        writer.writerow([id])
                        # Write data for table 1
                        writer.writerow([orientada[0][2]])
                        writer.writerow([orientada[0][1]])

                        # Add an empty line between tables
                        writer.writerow([])

                        # Add an empty line between tables
                        writer.writerow(["Secreta",orientada[0][-2]])

                    
                        # Add an empty line between tables
                        writer.writerow([])
                        writer.writerow(["Votos"])
                        writer.writerow(["Votos Sim",orientada[0][3]])
                        writer.writerow(["Votos Não",orientada[0][4]])
                        writer.writerow(["Abstenção",orientada[0][6]])
                        writer.writerow(["Não Compareceu",ncom])
                        writer.writerow(["Votos do Presidente",orientada[0][5]])
                        writer.writerow(["PNRV",orientada[0][7]])
                        writer.writerow(["Total",orientada[0][8]])

                        # Add an empty line between tables
                        writer.writerow([])
                        
                        writer.writerow(["Resultado",orientada[0][-3]])
                        writer.writerow(["Orientação do Governo",orientada[0][-1]])

                        # Add an empty line between tables
                        writer.writerow([])

                        # Write data for table 2
                        writer.writerow(['Parlamentar','Partido','UF','Voto'])
                        writer.writerows(l)
                
                    lista.append(os.path.join(settings.MEDIA_ROOT,  "CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv"))
                    lista.append(os.path.join(settings.MEDIA_ROOT,  'word/Votação '+str(cont)+ 'Plenário '+ano[2]+"_"+ano[1]+"_"+ano[0]+'.docx'))
                    cont+=1
                    csvfile.close()
                    z.close()
            
                
                                

            else:
                    api_end_point2 = "https://legis.senado.leg.br/dadosabertos/materia/"+str(votos['Votacao']['CodigoMateria'])
                    joke = requests.get(api_end_point2)
                    xpars = xmltodict.parse(joke.text)
                    materia =  xpars['DetalheMateria']
                    sim = 0
                    nao = 0
                    presidente  = 0 
                    abss = 0 
                    pnrv=0
                    ncom =0
                    desc = ""
                    mtr= "Sem Ementa"
                    res = ""
                    secreta=""
                    d= ""
                    l = [] 
                    orientada = []
                    id = ""
                

                    if 'Materia' in materia:
                        materia =  xpars['DetalheMateria']['Materia']['DadosBasicosMateria']['EmentaMateria']
                        mtr = materia
                        desc= votos['Votacao']['DescricaoVotacao']
                        d = votos['Votacao']['DataSessao']

                        id = votos['Votacao']['DescricaoIdentificacaoMateria']
                        if votos['Votacao']['Resultado']=="A":
                            res ="Aprovada"
                        else:
                            res = "Reprovada"

                        if votos['Votacao']['Secreta']=="S":
                            secreta ="Sim"
                        else:
                            secreta = "Não"

                        for j in votos['Votacao']['Votos']['VotoParlamentar']:
                        
                            if j['Voto']=='Sim':
                                sim+=1
                            elif j['Voto']=='Não':
                                nao+=1
                            elif j['Voto']=='P-NRV':
                                pnrv+=1
                            elif j['Voto']=='Presidente (art. 51 RISF)':
                                presidente+=1
                            elif  j['Voto']=='Abstenção':
                                abss+=1
                            elif j['Voto']=='NCom':
                                ncom+=1
                        

                        
                            l.append([j['NomeParlamentar'],j['SiglaPartido'],j['SiglaUF'],j['Voto']])


                     
                   

                    else:

                        desc= votos['Votacao']['DescricaoVotacao']
                        d = votos['Votacao']['DataSessao']
                        id = votos['Votacao']['DescricaoIdentificacaoMateria']
                        if votos['Votacao']['Resultado']=="A":
                            res ="Aprovada"
                        else:
                            res = "Reprovada"

                        if votos['Votacao']['Secreta']=="S":
                            secreta ="Sim"
                        else:
                            secreta = "Não"
                    
                        for j in votos['Votacao']['Votos']['VotoParlamentar']:
                        
                            if j['Voto']=='Sim':
                                sim+=1
                            elif j['Voto']=='Não':
                                nao+=1
                            elif j['Voto']=='P-NRV':
                                pnrv+=1
                            elif j['Voto']=='Presidente (art. 51 RISF)':
                                presidente+=1
                            elif  j['Voto']=='Abstenção':
                                abss+=1
                            elif j['Voto']=='NCom':
                                ncom+=1
                
                       
                            l.append([j['NomeParlamentar'],j['SiglaPartido'],j['SiglaUF'],j['Voto']])

                    api_end_point3 = "https://legis.senado.leg.br/dadosabertos/plenario/votacao/orientacaoBancada/"+action
                    joke = requests.get(api_end_point3)
                    xpars = xmltodict.parse(joke.text)
                    ori =  xpars['OrientacaoBancada']
                    
                    if ori is None:
                        return render(request,'plenario.html')

                    if 'votacoes' in ori:
                        ori = xpars['OrientacaoBancada']['votacoes']
                        flag = 0
                        
                            
                        if 'orientacoesLideranca' in ori:
                                
                            
                            for j in ori['orientacoesLideranca']:
                                    z = j['partido']
                                    if z=="Governo":
                                        print("teste")
                                        flag =1
                                        total = sim+nao+presidente+abss+pnrv+ncom
                                        orientada.append([d,desc,mtr,sim,nao,presidente,abss,pnrv,total,res,secreta, j['voto']])
                                        break
                                    if flag==1:
                                        break
                        if flag==0:
                            total = sim+nao+presidente+abss+pnrv+ncom
                            orientada.append([d,desc,mtr,sim,nao,presidente,abss,pnrv,total,res,secreta,"SEM ORIENTAÇÃO"])


                    print(cont)
                    document = Document()
                    font = document.styles['Normal'].font
                    font.name = 'Arial'
                    font.size = Pt(12)
                    styles = document.styles['Heading 1'].font
                    styles.name = 'Arial'
                    paragraph_format = document.styles['Normal'].paragraph_format
                    paragraph_format.line_spacing = 1.15
                    z = open("CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv",'w+',encoding='iso-8859-1')
               

                    document.add_heading('DataSessao ' + str(orientada[0][0]) +"\n", 0)

                    document.add_heading(id, 0)

                    document.add_paragraph("\n"+orientada[0][2])

                    document.add_paragraph("\n"+orientada[0][1])

                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Secreta'
                    hdr_cells[1].text = orientada[0][-2]

               
                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos'

                
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos Sim'
                    hdr_cells[1].text = str(orientada[0][3])

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos Não'
                    hdr_cells[1].text = str(orientada[0][4])
                
                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Abstenção'
                    hdr_cells[1].text = str(orientada[0][6])


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Votos do Presidente'
                    hdr_cells[1].text = str(orientada[0][5])


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Não Compareceu'
                    hdr_cells[1].text = str(ncom)


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'PNRV'
                    hdr_cells[1].text = str(orientada[0][7])

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Total'
                    hdr_cells[1].text = str(orientada[0][8])


                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Resultado'
                    hdr_cells[1].text = orientada[0][-3]


                    table = document.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Orientação do Governo'
                    hdr_cells[1].text = orientada[0][-1]


                    document.add_paragraph("\n")

                    table = document.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Parlamentar'
                    hdr_cells[1].text = 'Partido'
                    hdr_cells[2].text = 'UF'
                    hdr_cells[3].text = 'Voto'

                
                    for iii in l:
                        table = document.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = iii[0]
                        hdr_cells[1].text =  iii[1]
                        hdr_cells[2].text = iii[2]
                        hdr_cells[3].text =  iii[3]
                



                    document.save('word/Votação '+str(cont)+ 'Plenário '+ano[2]+"_"+ano[1]+"_"+ano[0]+'.docx')
                    print(orientada)   
                    l.sort(key=lambda x:x[1])     
                    with open("CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv",'w', newline='',encoding='cp1252') as csvfile:
                        writer = csv.writer(csvfile,delimiter =';')

                        # Write data for table 1
                        writer.writerow(["DataSessao",orientada[0][0]])
                      

                        # Add an empty line between tables
                        writer.writerow([])
                        
                        writer.writerow([id])
                        # Write data for table 1
                        writer.writerow([orientada[0][2]])
                        writer.writerow([orientada[0][1]])

                        # Add an empty line between tables
                        writer.writerow([])

                        # Add an empty line between tables
                        writer.writerow(["Secreta",orientada[0][-2]])

                    
                        # Add an empty line between tables
                        writer.writerow([])
                        writer.writerow(["Votos"])
                        writer.writerow(["Votos Sim",orientada[0][3]])
                        writer.writerow(["Votos Não",orientada[0][4]])
                        writer.writerow(["Abstenção",orientada[0][6]])
                        writer.writerow(["Não Compareceu",ncom])
                        writer.writerow(["Votos do Presidente",orientada[0][5]])
                        writer.writerow(["PNRV",orientada[0][7]])
                        writer.writerow(["Total",orientada[0][8]])

                        # Add an empty line between tables
                        writer.writerow([])
                        
                        writer.writerow(["Resultado",orientada[0][-3]])
                        writer.writerow(["Orientação do Governo",orientada[0][-1]])

                        # Add an empty line between tables
                        writer.writerow([])

                        # Write data for table 2
                        writer.writerow(['Parlamentar','Partido','UF','Voto'])
                        writer.writerows(l)
                
                    lista.append(os.path.join(settings.MEDIA_ROOT,  "CSV/Votação"+str(cont)+"_"+"PLEN"+"_"+ano[2]+"_"+ano[1]+"_"+ano[0]+ ".csv"))
                    lista.append(os.path.join(settings.MEDIA_ROOT,  'word/Votação '+str(cont)+ 'Plenário '+ano[2]+"_"+ano[1]+"_"+ano[0]+'.docx'))
                    cont+=1
                    csvfile.close()
                    z.close()
            
                
        if lista!=[]:    

            return  download_zip(lista,"Plenario_"+ano[2]+"_"+ano[1]+"_"+ano[0])             


            
            
            
                
    return render(request,'plenario.html',{'alerta': "falso"})


def comissao(request):

    lista=[]
    comi = "https://legis.senado.leg.br/dadosabertos/dados/ComissoesPermanentes.xml"
    joke = requests.get(comi)
    xpars = xmltodict.parse(joke.text)
    votos = xpars['ComissoesPermanentes']['Colegiados']['Colegiado']

    comis = []

    for i in votos:
        comis.append(i['SiglaColegiado'])

    if request.method=="POST":
       
        data = request.POST
        nomecomiss = data.get("comiss")
        ini = data.get("data1")
        fim = data.get("data2")

        ano1 = ini.split("-")
        ano2 = fim.split("-")
        
        ini = ini.replace("-","")
        fim = fim.replace("-","")
        cont=1    
        print(ini,fim,nomecomiss)

        api_end_point = "https://legis.senado.leg.br/dadosabertos/votacaoComissao/comissao/"+nomecomiss+"?dataInicio="+ini+"&"+"dataFim="+fim
        joke = requests.get(api_end_point)
        xpars = xmltodict.parse(joke.text)
        votos = xpars['VotacoesComissao']

        if "Votacoes" not in votos:
            print("Não teve votação nesse intervalo")
            return render(request,'comissao.html',{'comissao': comis,'alerta': "verdadeiro"})

        else:
            votos = xpars['VotacoesComissao']['Votacoes']
            
            DataHoraInicioReuniao = ""
            NumeroReuniaoColegiado = ""
            TipoReuniao = ""
            NomeColegiado = ""
            IdentificacaoMateria = ""
            DescricaoIdentificacaoMateria =""
            DescricaoVotacao = ""
            
            
            if type(votos['Votacao'])==list:

                for i in range(len(votos['Votacao'])):
               
             
                    sim = 0
                    nao = 0
                    presidente  = 0 
                    abss = 0 
                    pnrv=0
                    total = 0
                    header = []
                    DataHoraInicioReuniao = votos['Votacao'][i]['DataHoraInicioReuniao']
                    NumeroReuniaoColegiado = votos['Votacao'][i]['NumeroReuniaoColegiado']
                    TipoReuniao = votos['Votacao'][i]['TipoReuniao']
                    NomeColegiado =  votos['Votacao'][i]['NomeColegiado']
                    IdentificacaoMateria =  votos['Votacao'][i]['IdentificacaoMateria']
                    DescricaoIdentificacaoMateria = votos['Votacao'][i]['DescricaoIdentificacaoMateria']
                    DescricaoVotacao = votos['Votacao'][i]['DescricaoVotacao']

                    t = votos['Votacao'][i]['IdentificacaoMateria']
                    t = t.split(" ")

                    sigla = t[0]

                    r = t[1].split('/')
                
                    codigo = r[0]
                    ano = r[1]


                    api_end_point2 = "https://legis.senado.leg.br/dadosabertos/materia/"+sigla+"/"+codigo+"/"+ano
                    joke = requests.get(api_end_point2)
                    xpars = xmltodict.parse(joke.text)
                    materia = xpars['DetalheMateria']

                    ementa = "Sem ementa"

                    if "Materia" in  materia:
      
                        ementa = materia['Materia']['DadosBasicosMateria']['EmentaMateria']



                    body = []
                

                    for j in votos['Votacao'][i]['Votos']['Voto']:
                    
                        body.append([j['NomeParlamentar'],j['SiglaPartidoParlamentar'],j['QualidadeVoto']])
                    horas = DataHoraInicioReuniao.split('T')
                    header.append([horas[0],NumeroReuniaoColegiado,TipoReuniao,NomeColegiado,IdentificacaoMateria,DescricaoIdentificacaoMateria,DescricaoVotacao,votos['Votacao'][i]['TotalVotosSim'],votos['Votacao'][i]['TotalVotosNao'],votos['Votacao'][i]['TotalVotosAbstencao'],int(votos['Votacao'][i]['TotalVotosSim'])+int(votos['Votacao'][i]['TotalVotosNao'])+int(votos['Votacao'][i]['TotalVotosAbstencao'])])
    

               
                    z = open("Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv",'w+',encoding='iso-8859-1')
                    body.sort(key=lambda x:x[1])    

                    with open("Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv",'w', newline='',encoding='iso-8859-1') as csvfile:
                            writer = csv.writer(csvfile,delimiter =';')

                            # Write data for table 1
                            writer.writerow([header[0][3]])
                            writer.writerow([header[0][0]])

                            # Add an empty line between tables
                            writer.writerow([])


                      
                            writer.writerow([header[0][5]])
                      
                            writer.writerow([ementa])
                       
                            writer.writerow([header[0][6]])

                            # Add an empty line between tables
                            writer.writerow([])

                        
                            writer.writerow(['Votos'])
                            writer.writerow(['Votos Sim',header[0][7]])
                            writer.writerow(['Votos Não',header[0][8]])
                            writer.writerow(['Abstenção',header[0][9]])
                            writer.writerow(['Total',header[0][10]])

                            # Add an empty line between tables
                            writer.writerow([])

                            # Write data for table 2
                            writer.writerow(['Parlamentar','Partido','Voto'])
                            writer.writerows(body)
                    lista.append(os.path.join(settings.MEDIA_ROOT,  "Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv"))
                    cont+=1
                    csvfile.close()
                    z.close()

            else:
                    sim = 0
                    nao = 0
                    presidente  = 0 
                    abss = 0 
                    pnrv=0
                    total = 0
                    header = []
                    DataHoraInicioReuniao = votos['Votacao']['DataHoraInicioReuniao']
                    NumeroReuniaoColegiado = votos['Votacao']['NumeroReuniaoColegiado']
                    TipoReuniao = votos['Votacao']['TipoReuniao']
                    NomeColegiado =  votos['Votacao']['NomeColegiado']
                    if 'IdentificacaoMateria' in votos['Votacao']:
                        IdentificacaoMateria =  votos['Votacao']['IdentificacaoMateria']
                        DescricaoIdentificacaoMateria = votos['Votacao']['DescricaoIdentificacaoMateria']

                        t = votos['Votacao']['IdentificacaoMateria']
                        t = t.split(" ")

                        sigla = t[0]

                        r = t[1].split('/')
                
                        codigo = r[0]
                        ano = r[1]


                        api_end_point2 = "https://legis.senado.leg.br/dadosabertos/materia/"+sigla+"/"+codigo+"/"+ano
                        joke = requests.get(api_end_point2)
                        xpars = xmltodict.parse(joke.text)
                        materia = xpars['DetalheMateria']

                        ementa = "Sem ementa"

                        if "Materia" in  materia:
      
                            ementa = materia['Materia']['DadosBasicosMateria']['EmentaMateria']
                    else:
                        IdentificacaoMateria = "Não tem" 
                        DescricaoIdentificacaoMateria = "Não tem"
                        ementa = "Sem ementa"
                    DescricaoVotacao = votos['Votacao']['DescricaoVotacao']

                   



                    body = []
                

                    for j in votos['Votacao']['Votos']['Voto']:
                    
                        body.append([j['NomeParlamentar'],j['SiglaPartidoParlamentar'],j['QualidadeVoto']])
                    horas = DataHoraInicioReuniao.split('T')
                    header.append([horas[0],NumeroReuniaoColegiado,TipoReuniao,NomeColegiado,IdentificacaoMateria,DescricaoIdentificacaoMateria,DescricaoVotacao,votos['Votacao']['TotalVotosSim'],votos['Votacao']['TotalVotosNao'],votos['Votacao']['TotalVotosAbstencao'],int(votos['Votacao']['TotalVotosSim'])+int(votos['Votacao']['TotalVotosNao'])+int(votos['Votacao']['TotalVotosAbstencao'])])
    

               
                    z = open("Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv",'w+',encoding='iso-8859-1')
                    body.sort(key=lambda x:x[1])    

                    with open("Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv",'w', newline='',encoding='iso-8859-1') as csvfile:
                            writer = csv.writer(csvfile,delimiter =';')

                            # Write data for table 1
                            writer.writerow([header[0][3]])
                            writer.writerow([header[0][0]])

                            # Add an empty line between tables
                            writer.writerow([])


                      
                            writer.writerow([header[0][5]])
                      
                            writer.writerow([ementa])
                       
                            writer.writerow([header[0][6]])

                            # Add an empty line between tables
                            writer.writerow([])

                        
                            writer.writerow(['Votos'])
                            writer.writerow(['Votos Sim',header[0][7]])
                            writer.writerow(['Votos Não',header[0][8]])
                            writer.writerow(['Abstenção',header[0][9]])
                            writer.writerow(['Total',header[0][10]])

                            # Add an empty line between tables
                            writer.writerow([])

                            # Write data for table 2
                            writer.writerow(['Parlamentar','Partido','Voto'])
                            writer.writerows(body)
                    lista.append(os.path.join(settings.MEDIA_ROOT,  "Votação"+str(cont)+"_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0]+ ".csv"))
                    cont+=1
                    csvfile.close()
                    z.close()


        if lista!=[]:
            return  download_zip(lista,"Comissão_"+nomecomiss+"_"+ano1[2]+"_"+ano1[1]+"_"+ano1[0])
    
    return render(request,'comissao.html',{'comissao': comis,'alerta': "falso"})


def gerandoatabelataquigrafica():
    current_date = datetime.date.today()

    

    data = str(current_date).split('-')
    

    l = []

    for z in [data[0]]:
    
        i = int(data[1])
    
        if i < 10:
            apiendpoint = "https://legis.senado.leg.br/dadosabertos/agendareuniao/mes/"+z+"0"+str(i)+".json"
            joke = requests.get(apiendpoint)
        
            y = joke.json()
       
            if y['AgendaReuniao']['reunioes']!=None:
                print(i)
                for j in range(len(y['AgendaReuniao']['reunioes']['reuniao'])):
                    l.append(y['AgendaReuniao']['reunioes']['reuniao'][j]['codigo'])
        
        else:
            apiendpoint = "https://legis.senado.leg.br/dadosabertos/agendareuniao/mes/"+z+str(i)+".json"
            joke = requests.get(apiendpoint)
        
            y = joke.json()
       
            if y['AgendaReuniao']['reunioes']!=None:
                print(i)
                for j in range(len(y['AgendaReuniao']['reunioes']['reuniao'])):
                    l.append(y['AgendaReuniao']['reunioes']['reuniao'][j]['codigo'])

        texto = pd.read_csv('CSV/TaquiCompleta.csv',sep=";",encoding='cp1252', on_bad_lines='skip')
        texto = texto.values.tolist()



    l.sort()


    for i in l:
        apiendpoint = "https://legis.senado.gov.br/escriba-servicosweb/reuniao/xml/"+i
        joke = requests.get(apiendpoint)
    
        if joke.status_code ==200 :
     
            taqui = xmltodict.parse(joke.text)

            taquigraficas = taqui['notasTaquigraficas']['quartos']
        
            if type(taqui['notasTaquigraficas']['quartos']) == list:

                for j in range(len(taqui['notasTaquigraficas']['quartos'])):
                    if 'texto' in taqui['notasTaquigraficas']['quartos'][j]:
                        w = [taqui['notasTaquigraficas']['dadosReuniao']['nomesComissoes'],taqui['notasTaquigraficas']['dadosReuniao']['siglasComissoes'],taqui['notasTaquigraficas']['quartos'][j]['texto']]
                        if w not in texto:
                      
                            texto.append([taqui['notasTaquigraficas']['dadosReuniao']['nomesComissoes'],taqui['notasTaquigraficas']['dadosReuniao']['siglasComissoes'],taqui['notasTaquigraficas']['quartos'][j]['texto']])
                        
            else:
                w = [taqui['notasTaquigraficas']['dadosReuniao']['nomesComissoes'],taqui['notasTaquigraficas']['dadosReuniao']['siglasComissoes'],taqui['notasTaquigraficas']['quartos']['texto']]
                if w not in texto:
               
                    texto.append([taqui['notasTaquigraficas']['dadosReuniao']['nomesComissoes'],taqui['notasTaquigraficas']['dadosReuniao']['siglasComissoes'],taqui['notasTaquigraficas']['quartos']['texto']])

    texto = list(texto for texto,_ in itertools.groupby(texto))
    df1 = pd.DataFrame(texto,columns=['NomeComissao','sigla','fala'])
    df1.to_csv('CSV/TaquiCompleta.csv', sep=';', encoding='cp1252',index=False)

def taqui(request):
    lista=[]
    comi = "https://legis.senado.leg.br/dadosabertos/dados/ComissoesPermanentes.xml"
    joke = requests.get(comi)
    xpars = xmltodict.parse(joke.text)
    votos = xpars['ComissoesPermanentes']['Colegiados']['Colegiado']

    comis = []

    for i in votos:
        comis.append(i['SiglaColegiado'])
        gerandoatabelataquigrafica()
    return render(request,'taqui.html',{'comissao': comis,'alerta': "falso"})

